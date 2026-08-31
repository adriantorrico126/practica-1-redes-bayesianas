"""Crea el documento individual con el patrón de la práctica 2."""
from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).parent
OUTPUT = ROOT / "Patron_y_justificacion_Adrian_Torrico.docx"
COLAB_URL = "https://colab.research.google.com/github/adriantorrico126/practica-1-redes-bayesianas/blob/main/Practica_1_Redes_Bayesianas_Adrian_Torrico.ipynb"


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = OxmlElement("w:tcMar")
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = OxmlElement(f"w:{side}")
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
        tc_mar.append(node)
    tc_pr.append(tc_mar)


def shade(cell, color):
    shade_node = OxmlElement("w:shd")
    shade_node.set(qn("w:fill"), color)
    cell._tc.get_or_add_tcPr().append(shade_node)


def table(doc, headers, rows, widths):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.autofit = False
    for idx, header in enumerate(headers):
        c = t.rows[0].cells[idx]
        c.text = header
        c.width = Inches(widths[idx])
        shade(c, "E8EEF5")
        set_cell_margins(c)
        for run in c.paragraphs[0].runs:
            run.bold = True
    for row in rows:
        cells = t.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = str(value)
            cells[idx].width = Inches(widths[idx])
            set_cell_margins(cells[idx])
    return t


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    relation_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relation_id)
    run = OxmlElement("w:r")
    run_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    run_pr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_pr.append(underline)
    run.append(run_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


doc = Document()
section = doc.sections[0]
section.top_margin = section.bottom_margin = Inches(1)
section.left_margin = section.right_margin = Inches(1)
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.1
for name, size, color in (("Heading 1", 16, "2E74B5"), ("Heading 2", 13, "2E74B5")):
    style = doc.styles[name]
    style.font.name = "Calibri"
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.paragraph_format.space_before = Pt(14)
    style.paragraph_format.space_after = Pt(6)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(54)
r = p.add_run("PATRÓN DE MINERÍA DE DATOS")
r.bold = True
r.font.size = Pt(23)
r.font.color.rgb = RGBColor(31, 77, 120)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Práctica individual - Minería de Datos")
r.font.size = Pt(14)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("Estudiante: Adrian Torrico\nFecha: 31 de agosto de 2026")

doc.add_heading("1. Patrón propuesto", 1)
doc.add_paragraph("El patrón que identifiqué es el siguiente:")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Si un cliente tiene contrato mensual, no cuenta con soporte técnico y presenta satisfacción baja, entonces tiene una probabilidad elevada de abandonar el servicio.")
r.bold = True
r.font.size = Pt(12)
doc.add_paragraph("En notación de regla de asociación: {Contract_Type = Monthly, Technical_Support = No, Satisfaction = Low} → {Churn = Yes}.")
table(doc, ["Medida", "Cálculo", "Resultado"], [
    ["Soporte", "17 de 600 clientes cumplen el antecedente", "2.83%"],
    ["Confianza", "10 de los 17 clientes del grupo abandonaron", "58.82%"],
    ["Tasa base", "153 abandonos entre 600 clientes", "25.50%"],
    ["Lift", "58.82% / 25.50%", "2.31"],
], [1.35, 3.6, 1.15])

doc.add_heading("2. Por qué el patrón cumple las características solicitadas", 1)
doc.add_heading("Válido", 2)
doc.add_paragraph("Lo considero válido porque no surge de un caso aislado: el antecedente aparece en 17 clientes y 10 de ellos abandonaron. Su confianza de 58.82% supera la tasa base de abandono de 25.50%. El lift de 2.31 significa que el abandono es 2.31 veces más frecuente en este grupo que en el conjunto completo. Estas medidas respaldan una asociación descriptiva; no prueban causalidad.")
doc.add_heading("Novedoso", 2)
doc.add_paragraph("El patrón no se limita a observar una sola variable. Combina tres señales que, juntas, describen una situación específica de riesgo: compromiso contractual bajo, ausencia de ayuda ante problemas y percepción negativa del servicio. Esta combinación entrega una lectura más informativa que decir solamente que la satisfacción baja se relaciona con abandono.")
doc.add_heading("Útil", 2)
doc.add_paragraph("La regla permite priorizar acciones concretas de retención. Por ejemplo, la empresa puede identificar primero a los clientes mensuales sin soporte y con satisfacción baja, ofrecerles soporte proactivo y proponerles una alternativa de contrato. En el conjunto analizado, esto focaliza la atención en 17 clientes, de los cuales 10 abandonaron.")
doc.add_heading("Comprensible", 2)
doc.add_paragraph("La regla está escrita con atributos que una persona de negocio entiende sin conocimientos de estadística avanzada. Además, se interpreta como una frase directa: cliente mensual + sin soporte + insatisfecho = alto riesgo de abandono.")

doc.add_heading("3. Cómo encontré el patrón", 1)
doc.add_paragraph("Seguí un procedimiento que combinó intuición de un experto del negocio con evidencia de la práctica previa.")
table(doc, ["Etapa", "Evidencia o decisión"], [
    ["1. Exploración previa", "La tasa de abandono fue 31.31% con contrato mensual, 38.24% sin soporte técnico y 52.78% con satisfacción baja."],
    ["2. Intuición de experto", "Supuse que un cliente con poco compromiso contractual, sin ayuda técnica e insatisfecho tendría mayor predisposición a abandonar."],
    ["3. Combinación de señales", "Construí la tabla de contingencia para la intersección de las tres condiciones y contrasté sus abandonos con el total."],
    ["4. Evaluación", "Calculé soporte, confianza y lift. Elegí la regla porque mantiene 17 casos, evitando una combinación demasiado específica con pocos ejemplos."],
    ["5. Contraste con la red", "La red bayesiana previa estimó 52.82% de abandono para un perfil aún más específico que añade poca antigüedad, coherente con la dirección del patrón."],
], [1.35, 4.75])

doc.add_heading("4. Alcance y limitaciones", 1)
doc.add_paragraph("El patrón es útil para segmentación y prevención, pero no debe interpretarse como una prueba de que la ausencia de soporte o el contrato mensual causen directamente el abandono. También conviene validar la regla con nuevos periodos de datos antes de usarla en decisiones operativas. Para mejorarla, podría incorporarse el historial de reclamos, interrupciones del servicio y contactos previos de retención.")

doc.add_heading("5. Enlace al programa en Python (Google Colab)", 1)
p = doc.add_paragraph("El notebook reproducible, con la preparación de datos, las redes bayesianas y el análisis, está disponible en: ")
add_hyperlink(p, "Abrir programa en Google Colab", COLAB_URL)
doc.save(OUTPUT)
print(OUTPUT)
