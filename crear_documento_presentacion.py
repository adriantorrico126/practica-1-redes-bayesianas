"""Documento narrativo para presentar el patrón de abandono."""
from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).parent
OUT = ROOT / "Documento_presentacion_patron_Adrian_Torrico.docx"
COLAB = "https://colab.research.google.com/github/adriantorrico126/practica-1-redes-bayesianas/blob/main/Practica_1_Redes_Bayesianas_Adrian_Torrico.ipynb"


def set_cell_margins(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = OxmlElement("w:tcMar")
    for side, value in (("top", 80), ("start", 120), ("bottom", 80), ("end", 120)):
        node = OxmlElement(f"w:{side}")
        node.set(qn("w:w"), str(value)); node.set(qn("w:type"), "dxa"); tc_mar.append(node)
    tc_pr.append(tc_mar)


def shade(cell, color):
    n = OxmlElement("w:shd"); n.set(qn("w:fill"), color); cell._tc.get_or_add_tcPr().append(n)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers)); table.style = "Table Grid"; table.autofit = False
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]; cell.text = header; cell.width = Inches(widths[i]); shade(cell, "F4F6F9"); set_cell_margins(cell)
        for run in cell.paragraphs[0].runs: run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value); cells[i].width = Inches(widths[i]); set_cell_margins(cells[i])
    return table


def add_link(paragraph, text, url):
    rel = paragraph.part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    h = OxmlElement("w:hyperlink"); h.set(qn("r:id"), rel)
    r = OxmlElement("w:r"); rpr = OxmlElement("w:rPr")
    c = OxmlElement("w:color"); c.set(qn("w:val"), "0563C1"); rpr.append(c)
    u = OxmlElement("w:u"); u.set(qn("w:val"), "single"); rpr.append(u)
    r.append(rpr); t = OxmlElement("w:t"); t.text = text; r.append(t); h.append(r); paragraph._p.append(h)


doc = Document()
sec = doc.sections[0]
sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1)
sec.header_distance = sec.footer_distance = Inches(0.492)
normal = doc.styles["Normal"]
normal.font.name = "Calibri"; normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri"); normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(8); normal.paragraph_format.line_spacing = 1.333
for style_name, size, color, before, after in (("Heading 1", 16, "2E74B5", 18, 10), ("Heading 2", 13, "2E74B5", 12, 6), ("Heading 3", 12, "1F4D78", 8, 4)):
    st = doc.styles[style_name]; st.font.name = "Calibri"; st.font.size = Pt(size); st.font.color.rgb = RGBColor.from_string(color)
    st.paragraph_format.space_before = Pt(before); st.paragraph_format.space_after = Pt(after)
footer = sec.footer.paragraphs[0]; footer.alignment = WD_ALIGN_PARAGRAPH.CENTER; footer.add_run("Patrón de abandono de clientes | Adrian Torrico").font.size = Pt(9)

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before = Pt(68); p.paragraph_format.space_after = Pt(8)
r = p.add_run("PRESENTACIÓN DEL PATRÓN ENCONTRADO"); r.bold = True; r.font.size = Pt(24); r.font.color.rgb = RGBColor(31, 77, 120)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Abandono de clientes en servicios de Internet"); r.font.size = Pt(16); r.font.color.rgb = RGBColor(68, 68, 68)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.add_run("Minería de Datos\nAdrian Torrico\n31 de agosto de 2026")
doc.add_page_break()

doc.add_heading("1. Punto de partida", 1)
doc.add_paragraph("En este trabajo partí de una pregunta muy concreta: ¿qué características permiten reconocer a un cliente con mayor riesgo de abandonar el servicio de Internet? Mi objetivo no fue solamente obtener un porcentaje, sino traducir los datos en una regla clara que pudiera comprender y comunicar con facilidad.")
doc.add_paragraph("Trabajé con 600 clientes. En el conjunto completo, 153 abandonaron el servicio; por ello, la tasa general de abandono fue 25.50%. Este valor fue mi referencia: cualquier patrón interesante debía mostrar un riesgo claramente superior a ese promedio.")
doc.add_picture(str(ROOT / "resultados" / "eda_churn.png"), width=Inches(6.2))
p = doc.add_paragraph("Figura 1. Contexto del abandono y comportamiento según tipo de contrato."); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in p.runs: run.italic = True; run.font.size = Pt(9)

doc.add_heading("2. Lo que observé antes de proponer la regla", 1)
doc.add_paragraph("Primero revisé cada variable por separado. Esta exploración me permitió pasar de una impresión general a señales específicas. Encontré que el abandono era mayor entre clientes con contrato mensual, sin soporte técnico y con satisfacción baja.")
add_table(doc, ["Señal observada", "Tasa de abandono", "Lectura que hice"], [
    ["Contrato mensual", "31.31%", "El compromiso es menor que en un contrato de mayor duración."],
    ["Sin soporte técnico", "38.24%", "Un problema técnico puede quedar sin una vía clara de solución."],
    ["Satisfacción baja (1-4)", "52.78%", "La percepción negativa del servicio es una alerta directa."],
], [2.0, 1.35, 3.15])
doc.add_paragraph("Mi intuición fue que estas tres condiciones no debían analizarse de manera aislada. Un cliente mensual puede abandonar con facilidad; si además no recibe soporte y está insatisfecho, su situación es más frágil. Esta idea se convirtió en una hipótesis que luego validé con los datos.")

doc.add_heading("3. El patrón que identifiqué", 1)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Contrato mensual + sin soporte técnico + satisfacción baja → alto riesgo de abandono"); r.bold = True; r.font.size = Pt(13); r.font.color.rgb = RGBColor(31, 77, 120)
doc.add_paragraph("Al revisar la combinación, encontré 17 clientes que cumplían simultáneamente las tres condiciones. Diez de ellos abandonaron el servicio. Esta es la evidencia cuantitativa de la regla:")
add_table(doc, ["Indicador", "Resultado", "Qué significa"], [
    ["Soporte", "17 de 600 (2.83%)", "La regla aparece en un grupo real y no en un caso aislado."],
    ["Confianza", "10 de 17 (58.82%)", "Más de la mitad de este grupo abandonó."],
    ["Lift", "2.31", "El abandono es 2.31 veces más frecuente que en el total."],
], [1.25, 1.6, 3.65])
doc.add_paragraph("Para explicarlo de forma cotidiana, es parecido a detectar que ciertos clientes de una cafetería dejan de volver cuando tuvieron una mala experiencia, no recibieron atención y no tienen ningún compromiso con el lugar. La combinación no garantiza que todos se vayan, pero permite actuar antes de que se marchen.")

doc.add_heading("4. Por qué considero que este patrón es valioso", 1)
doc.add_heading("Es válido", 2)
doc.add_paragraph("Lo sostengo con soporte, confianza y lift. La confianza de 58.82% es superior a la tasa base de 25.50%, y el lift confirma que no se trata de un resultado esperado por azar dentro del promedio. Aun así, lo presento como una asociación descriptiva, no como una afirmación causal.")
doc.add_heading("Es novedoso", 2)
doc.add_paragraph("No me limité a una sola característica. La regla combina contrato, soporte y satisfacción para formar un perfil de riesgo. De esa manera, el resultado ofrece una lectura más específica que afirmar solamente que la satisfacción baja se relaciona con el abandono.")
doc.add_heading("Es útil y comprensible", 2)
doc.add_paragraph("Es útil porque indica dónde enfocaría una estrategia de retención: identificaría primero a este segmento y le ofrecería contacto proactivo, soporte técnico y una alternativa de contratación. Es comprensible porque puede resumirse en una frase que entiende un área comercial, de soporte o de gestión sin requerir conocimientos de redes bayesianas.")

doc.add_heading("5. Relación con el análisis previo", 1)
doc.add_paragraph("En la práctica anterior construí una Red Bayesiana propuesta y otra aprendida con Hill Climbing. La consulta sobre un cliente con contrato mensual, sin soporte, satisfacción baja y poca antigüedad produjo una probabilidad de abandono de 52.82%. Ese resultado fue coherente con el patrón actual y me ayudó a seleccionar las variables relevantes.")
doc.add_picture(str(ROOT / "resultados" / "eda_riesgos.png"), width=Inches(6.2))
p = doc.add_paragraph("Figura 2. El abandono aumenta de forma notoria cuando la satisfacción es baja."); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in p.runs: run.italic = True; run.font.size = Pt(9)
doc.add_paragraph("Sin embargo, evité usar como patrón principal la combinación que añadía poca antigüedad, porque solo aparecía en cinco clientes. Preferí una regla con tres condiciones y 17 casos: conserva una señal fuerte, pero tiene un soporte más razonable para ser defendida en la presentación.")

doc.add_heading("6. Cierre", 1)
doc.add_paragraph("Concluyo que los datos no solo permitieron estimar probabilidades: también permitieron construir una regla accionable. El patrón identifica un grupo pequeño pero relevante, con una probabilidad de abandono de 58.82%. Si yo estuviera apoyando a la empresa, usaría esta regla como un criterio inicial de priorización y luego validaría sus resultados con nuevos periodos de datos.")
p = doc.add_paragraph("El programa reproducible está disponible en: "); add_link(p, "Google Colab", COLAB)
doc.save(OUT)
print(OUT)
