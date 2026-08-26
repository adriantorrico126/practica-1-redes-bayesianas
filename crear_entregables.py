"""Crea el notebook de Colab y el informe DOCX a partir del analisis ejecutado."""
import json
from pathlib import Path

import nbformat as nbf
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).parent
OUT = ROOT / "resultados"
with open(OUT / "resumen.json", encoding="utf-8") as f:
    R = json.load(f)


def create_notebook():
    nb = nbf.v4.new_notebook()
    cells = []
    def md(text): cells.append(nbf.v4.new_markdown_cell(text))
    def code(text): cells.append(nbf.v4.new_code_cell(text))
    md("# Práctica 1: Redes Bayesianas para predecir abandono de clientes\n\n**Estudiante:** Adrian Torrico  \\n**Materia:** Minería de Datos\n\nEn este notebook desarrollo y evalúo dos Redes Bayesianas discretas para estimar la probabilidad de abandono (churn) de clientes de Internet.")
    md("## 0. Preparación\n\nSubo `customers (2).csv` al entorno de Colab antes de ejecutar las celdas. La versión fijada de `pgmpy` mantiene la compatibilidad del código y hace el resultado reproducible.")
    code("!pip install -q pgmpy==0.1.25\n\nimport pandas as pd\nimport numpy as np\nimport matplotlib.pyplot as plt\nimport seaborn as sns\nimport networkx as nx\nfrom pgmpy.models import BayesianNetwork\nfrom pgmpy.estimators import BayesianEstimator, HillClimbSearch, BicScore\nfrom pgmpy.inference import VariableElimination\nfrom sklearn.model_selection import train_test_split\nfrom sklearn.metrics import accuracy_score, precision_score, recall_score, f1_score, confusion_matrix, classification_report\n\nRANDOM_STATE = 42\nsns.set_theme(style='whitegrid')")
    md("## Actividad 1. Exploración de los datos")
    code("df = pd.read_csv('customers (2).csv')\nprint('Dimensiones:', df.shape)\ndisplay(df.head())\ndisplay(df.info())\ndisplay(df.isna().sum().to_frame('Valores faltantes'))\ndisplay(df.describe().T)\n\nfig, axes = plt.subplots(1, 2, figsize=(12,4))\nsns.countplot(data=df, x='Churn', hue='Churn', legend=False, ax=axes[0], palette={'No':'#3B82F6','Yes':'#EF4444'})\naxes[0].set_title('Distribución de abandono')\nsns.countplot(data=df, x='Contract_Type', hue='Churn', ax=axes[1])\naxes[1].set_title('Abandono según tipo de contrato')\nplt.tight_layout()\nplt.show()")
    md("Interpretación: reviso el balance de la variable objetivo, los faltantes y el comportamiento de las variables. `Internet_Service` tiene valores faltantes; los conservo como una categoría explícita (`Unknown`) para no eliminar clientes ni inventar información.")
    md("## Actividad 2. Preparación y discretización")
    code("def discretize(df):\n    d = df.copy()\n    d['Internet_Service'] = d['Internet_Service'].fillna('Unknown')\n    d['Age_Group'] = pd.cut(d['Age'], [17,30,45,60,70], labels=['Young','Adult','Mature','Senior'], include_lowest=True).astype(str)\n    d['Tenure_Group'] = pd.cut(d['Tenure_Months'], [-1,12,24,48,np.inf], labels=['Low','Medium','Established','Loyal'], include_lowest=True).astype(str)\n    d['Amount_Group'] = pd.qcut(d['Monthly_Amount'], 3, labels=['Low','Medium','High'], duplicates='drop').astype(str)\n    d['Complaints_Group'] = pd.cut(d['Complaints'], [-1,0,2,np.inf], labels=['None','Few','Many'], include_lowest=True).astype(str)\n    d['Satisfaction_Group'] = pd.cut(d['Satisfaction'], [0,4,7,10], labels=['Low','Medium','High'], include_lowest=True).astype(str)\n    cols = ['Age_Group','Tenure_Group','Amount_Group','Contract_Type','Payment_Method','Technical_Support','Internet_Service','Complaints_Group','Satisfaction_Group','Churn']\n    return d[cols].astype(str)\n\ndata = discretize(df)\ndisplay(pd.DataFrame({'Variable':data.columns, 'Categorías':[', '.join(sorted(data[c].unique())) for c in data.columns]}))\n\nfig, axes = plt.subplots(1, 2, figsize=(12,4))\nfor ax, col, order in [(axes[0], 'Satisfaction_Group', ['Low','Medium','High']), (axes[1], 'Tenure_Group', ['Low','Medium','Established','Loyal'])]:\n    rate = pd.crosstab(data[col], data['Churn'], normalize='index')['Yes'].mul(100).reindex(order)\n    rate.plot(kind='bar', ax=ax, color='#EF4444')\n    ax.set_title(f'Tasa de abandono por {col}')\n    ax.set_ylabel('Abandono (%)'); ax.set_xlabel('')\n    ax.tick_params(axis='x', rotation=0)\nplt.tight_layout(); plt.show()")
    md("## División de entrenamiento y prueba")
    code("train, test = train_test_split(data, test_size=0.25, stratify=data['Churn'], random_state=RANDOM_STATE)\nprint(f'Entrenamiento: {len(train)} clientes | Prueba: {len(test)} clientes')")
    md("## Actividad 3. Red Bayesiana propuesta")
    md("Propongo que contrato, soporte, satisfacción, antigüedad, reclamos y servicio de Internet influyan directamente en el abandono. También conecto edad, monto y reclamos con satisfacción, y contrato con antigüedad. Es una estructura explicable basada en el proceso de negocio.")
    code("manual_edges = [('Age_Group','Satisfaction_Group'), ('Amount_Group','Satisfaction_Group'), ('Complaints_Group','Satisfaction_Group'), ('Contract_Type','Tenure_Group'), ('Contract_Type','Churn'), ('Technical_Support','Churn'), ('Satisfaction_Group','Churn'), ('Tenure_Group','Churn'), ('Complaints_Group','Churn'), ('Internet_Service','Churn')]\nmanual = BayesianNetwork(manual_edges)\nmanual.add_nodes_from(data.columns)\nmanual.fit(train, estimator=BayesianEstimator, prior_type='BDeu', equivalent_sample_size=10)\n\ndef draw_network(model, title):\n    g = nx.DiGraph(model.edges())\n    plt.figure(figsize=(10,6))\n    pos = nx.spring_layout(g, seed=RANDOM_STATE, k=1.3)\n    nx.draw(g, pos, with_labels=True, node_color=['#FCA5A5' if n=='Churn' else '#BFDBFE' for n in g.nodes()], node_size=2200, arrowsize=20, font_size=8, font_weight='bold')\n    plt.title(title); plt.show()\n\ndraw_network(manual, 'Red Bayesiana propuesta')\nfor cpd in manual.get_cpds():\n    print(cpd)\n    print('-'*70)")
    md("Las salidas anteriores son las tablas de probabilidad condicional (CPD) estimadas con suavizado BDeu. El suavizado evita probabilidades exactamente cero para combinaciones poco frecuentes.")
    md("## Actividad 4. Estructura automática con Hill Climbing")
    code("hc = HillClimbSearch(train)\nlearned_structure = hc.estimate(scoring_method=BicScore(train), max_indegree=3, max_iter=1000, show_progress=False)\nautomatic = BayesianNetwork(learned_structure.edges())\nautomatic.add_nodes_from(data.columns)\nautomatic.fit(train, estimator=BayesianEstimator, prior_type='BDeu', equivalent_sample_size=10)\n\nprint('Aristas encontradas:')\nprint(list(automatic.edges()))\ndraw_network(automatic, 'Red Bayesiana aprendida con Hill Climbing')\nfor cpd in automatic.get_cpds():\n    print(cpd)\n    print('-'*70)")
    md("## Actividad 5. Inferencia probabilística")
    code("inference = VariableElimination(automatic)\nevidence = {'Contract_Type':'Monthly', 'Technical_Support':'No', 'Satisfaction_Group':'Low', 'Tenure_Group':'Low'}\nquery = inference.query(variables=['Churn'], evidence=evidence, show_progress=False)\nprint('Evidencia:', evidence)\nprint(query)\nprint(f\"Probabilidad de abandono: {query.values[list(query.state_names['Churn']).index('Yes')]:.2%}\")")
    md("## Actividad 6. Evaluación de los modelos")
    code("def evaluate(model, test):\n    features = [c for c in test.columns if c != 'Churn']\n    pred = model.predict(test[features])['Churn']\n    y = test['Churn']\n    return {'Accuracy':accuracy_score(y,pred), 'Precision':precision_score(y,pred,pos_label='Yes',zero_division=0), 'Recall':recall_score(y,pred,pos_label='Yes',zero_division=0), 'F1-score':f1_score(y,pred,pos_label='Yes',zero_division=0)}, confusion_matrix(y,pred,labels=['No','Yes']), classification_report(y,pred,zero_division=0)\n\nmanual_metrics, manual_cm, manual_report = evaluate(manual, test)\nauto_metrics, auto_cm, auto_report = evaluate(automatic, test)\nmetrics = pd.DataFrame([manual_metrics, auto_metrics], index=['Red propuesta','Hill Climbing']).round(4)\ndisplay(metrics)\nprint('Matriz de confusión — Red propuesta [No, Yes]:\\n', manual_cm)\nprint('Matriz de confusión — Hill Climbing [No, Yes]:\\n', auto_cm)\nprint(auto_report)\nmetrics.plot(kind='bar', figsize=(10,4), ylim=(0,1), rot=0)\nplt.title('Comparación de métricas'); plt.ylabel('Valor'); plt.tight_layout(); plt.show()")
    md("## Actividad 7. Comparación y conclusión\n\nLa red propuesta usa conocimiento del problema y por ello sus relaciones son fáciles de justificar y comunicar. La red de Hill Climbing explora los datos y puede descubrir asociaciones no anticipadas, pero una arista no debe interpretarse automáticamente como causalidad. Comparo las métricas en el conjunto de prueba y priorizo Recall/F1 para churn, porque Accuracy puede verse favorecida por la clase mayoritaria de clientes que no abandonan. Una mejora posterior sería ajustar el umbral de decisión, incorporar más datos y validar mediante particiones repetidas.")
    nb['cells'] = cells
    nb['metadata'] = {"colab": {"name": "Practica_1_Redes_Bayesianas_Adrian_Torrico.ipynb", "provenance": []}, "kernelspec": {"name": "python3", "display_name": "Python 3"}, "language_info": {"name": "python"}}
    with open(ROOT / "Practica_1_Redes_Bayesianas_Adrian_Torrico.ipynb", "w", encoding="utf-8") as f: nbf.write(nb, f)


def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr(); shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), fill); tcPr.append(shd)

def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr(); tcMar = tcPr.first_child_found_in('w:tcMar')
    if tcMar is None: tcMar = OxmlElement('w:tcMar'); tcPr.append(tcMar)
    for m, v in [('top',top),('start',start),('bottom',bottom),('end',end)]:
        node = tcMar.find(qn(f'w:{m}'))
        if node is None: node = OxmlElement(f'w:{m}'); tcMar.append(node)
        node.set(qn('w:w'), str(v)); node.set(qn('w:type'), 'dxa')

def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers)); table.style = 'Table Grid'; table.autofit = False
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]; cell.text = str(h); shade(cell, 'E8EEF5'); set_cell_margins(cell)
        for run in cell.paragraphs[0].runs: run.bold = True
        if widths: cell.width = Inches(widths[i])
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value); set_cell_margins(cells[i])
            if widths: cells[i].width = Inches(widths[i])
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table

def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet'); p.add_run(text); return p

def add_figure(doc, file, caption, width=6.2):
    doc.add_picture(str(OUT / file), width=Inches(width))
    p = doc.add_paragraph(caption); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs: run.italic = True; run.font.size = Pt(9)

def create_docx():
    doc = Document()
    sec = doc.sections[0]; sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1)
    sec.header_distance = sec.footer_distance = Inches(0.492)
    styles = doc.styles
    styles['Normal'].font.name = 'Calibri'; styles['Normal']._element.rPr.rFonts.set(qn('w:ascii'), 'Calibri'); styles['Normal'].font.size = Pt(11)
    styles['Normal'].paragraph_format.space_after = Pt(6); styles['Normal'].paragraph_format.line_spacing = 1.1
    for name, size, color in [('Heading 1',16,'2E74B5'),('Heading 2',13,'2E74B5'),('Heading 3',12,'1F4D78')]:
        st=styles[name]; st.font.name='Calibri'; st.font.size=Pt(size); st.font.color.rgb=RGBColor.from_string(color); st.paragraph_format.space_before=Pt(14 if name=='Heading 1' else 10); st.paragraph_format.space_after=Pt(6)
    footer = sec.footer.paragraphs[0]; footer.alignment = WD_ALIGN_PARAGRAPH.CENTER; footer.add_run('Práctica 1 — Redes Bayesianas | Adrian Torrico').font.size = Pt(9)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(72); p.paragraph_format.space_after=Pt(10)
    r=p.add_run('PRÁCTICA 1: REDES BAYESIANAS'); r.bold=True; r.font.size=Pt(24); r.font.color.rgb=RGBColor(31,77,120)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run('Estimación de abandono de clientes de Internet'); r.font.size=Pt(16); r.font.color.rgb=RGBColor(68,68,68)
    doc.add_paragraph('Informe de resolución', style='Subtitle').alignment=WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('\nEstudiante: Adrian Torrico\nMateria: Minería de Datos\nFecha: 26 de agosto de 2026').alignment=WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()
    doc.add_heading('Resumen ejecutivo', 1)
    doc.add_paragraph(f'En esta práctica construí dos Redes Bayesianas discretas para estimar el abandono de clientes de una empresa de Internet. Trabajé con {R["n_rows"]} registros y separé los datos de manera estratificada en {R["train_rows"]} observaciones para entrenamiento y {R["test_rows"]} para prueba. La tasa observada de abandono fue de {R["churn_rate"]:.1%}.')
    doc.add_paragraph(f'Para el perfil solicitado (contrato mensual, sin soporte técnico, satisfacción baja y poca antigüedad), la red aprendida con Hill Climbing estima una probabilidad de abandono de {R["inference_probabilities"]["Yes"]:.2%}. Esta cifra supera ampliamente la tasa base y por ello identifica un perfil que requiere intervención prioritaria.')
    doc.add_heading('1. Objetivo y datos', 1)
    doc.add_paragraph('Mi objetivo fue modelar la probabilidad de que un cliente abandone el servicio (Churn = Yes) mediante Redes Bayesianas. Este enfoque representa dependencias probabilísticas entre características del cliente y permite responder consultas condicionadas a evidencia específica.')
    add_table(doc, ['Elemento','Resultado'], [['Registros',R['n_rows']],['Variables originales',R['n_columns']],['Variable objetivo','Churn (Yes/No)'],['Valores faltantes',f"Internet_Service: {R['missing_internet_service']}"],['Abandono observado',f"{R['churn_rate']:.1%} ({R['churn_counts']['Yes']} clientes)"]], [2.2,4.3])
    doc.add_heading('2. Actividad 1 — Exploración de datos', 1)
    doc.add_paragraph('Primero verifiqué las dimensiones, tipos de datos, valores faltantes y distribución de la variable objetivo. Detecté valores faltantes únicamente en Internet_Service; en lugar de descartar esos registros, los representé como la categoría Unknown. Esta decisión conserva la información disponible y permite que la red aprenda si la ausencia del dato aporta señal.')
    add_figure(doc, 'eda_churn.png', 'Figura 1. Distribución de abandono y tasa de abandono por tipo de contrato.')
    add_figure(doc, 'eda_riesgos.png', 'Figura 2. Riesgo de abandono según satisfacción y antigüedad.')
    doc.add_paragraph('Los patrones descriptivos indican que la satisfacción y la antigüedad merecen especial atención. A continuación reporto las tasas de abandono calculadas por grupos relevantes:')
    rr=R['risk_rates']; add_table(doc, ['Variable','Categoría','Abandono (%)'], [[ 'Contrato',k,f'{v:.2f}'] for k,v in rr['contract'].items()] + [[ 'Soporte técnico',k,f'{v:.2f}'] for k,v in rr['support'].items()] + [[ 'Satisfacción',k,f'{v:.2f}'] for k,v in rr['satisfaction'].items()] + [[ 'Antigüedad',k,f'{v:.2f}'] for k,v in rr['tenure'].items()], [1.7,2.6,2.2])
    doc.add_heading('3. Actividad 2 — Preparación de datos', 1)
    doc.add_paragraph('Como una Red Bayesiana discreta requiere estados categóricos, convertí las variables numéricas en intervalos interpretables. Elegí puntos de corte que conservan sentido de negocio y, para el monto mensual, tres grupos de tamaño aproximado similar mediante cuantiles.')
    add_table(doc, ['Variable original','Categorías aplicadas'], [['Age','18–30: Young; 31–45: Adult; 46–60: Mature; 61–70: Senior'],['Tenure_Months','0–12: Low; 13–24: Medium; 25–48: Established; 49+: Loyal'],['Monthly_Amount','Low, Medium y High (terciles por cuantiles)'],['Complaints','0: None; 1–2: Few; 3+: Many'],['Satisfaction','1–4: Low; 5–7: Medium; 8–10: High']], [2.0,4.5])
    doc.add_heading('4. Actividad 3 — Red Bayesiana propuesta', 1)
    doc.add_paragraph('Definí una estructura guiada por conocimiento del problema. Consideré que el contrato, soporte técnico, satisfacción, antigüedad, reclamos y tipo de servicio pueden influir directamente en Churn. Además, conecté reclamos, edad y monto con satisfacción; y el tipo de contrato con antigüedad. Estas relaciones no afirman causalidad demostrada: son hipótesis de modelado transparentes que pueden ser justificadas y discutidas.')
    add_figure(doc, 'red_manual.png', 'Figura 3. Estructura propuesta de la Red Bayesiana.')
    doc.add_paragraph('Estimé las tablas de probabilidad condicional mediante el estimador bayesiano BDeu (tamaño de muestra equivalente = 10). El notebook adjunto imprime las CPD completas de cada nodo, por lo que las tablas son reproducibles y verificables al ejecutar las celdas.')
    doc.add_heading('5. Actividad 4 — Red aprendida con Hill Climbing', 1)
    doc.add_paragraph('Apliqué Hill Climbing con BIC como función de puntuación y un máximo de tres padres por nodo. El algoritmo busca iterativamente la estructura que mejora la puntuación BIC; así equilibra ajuste a los datos y complejidad de la red.')
    add_figure(doc, 'red_hillclimb.png', 'Figura 4. Estructura descubierta automáticamente con Hill Climbing.')
    add_table(doc, ['Arista aprendida','Lectura descriptiva'], [[f'{a} → {b}','Relación seleccionada por el criterio BIC'] for a,b in R['automatic_edges']], [3.1,3.4])
    doc.add_paragraph('Las aristas aprendidas son asociaciones estadísticas condicionadas al conjunto de datos y a la función de puntuación; por tanto, no las interpreto como prueba de causalidad. Las tablas CPD de esta segunda red también se generan y muestran automáticamente en el notebook.')
    doc.add_heading('6. Actividad 5 — Inferencia probabilística', 1)
    ev=', '.join([f'{k} = {v}' for k,v in R['inference_evidence'].items()])
    doc.add_paragraph(f'Consulté la red aprendida usando la evidencia: {ev}.')
    add_table(doc, ['Resultado de la consulta','Probabilidad'], [['Churn = No',f"{R['inference_probabilities']['No']:.2%}"],['Churn = Yes',f"{R['inference_probabilities']['Yes']:.2%}"]], [3.4,3.1])
    doc.add_paragraph(f'Concluyo que este perfil tiene una probabilidad estimada de abandono de {R["inference_probabilities"]["Yes"]:.2%}. Recomendaría una acción de retención combinando contacto proactivo, mejora del soporte técnico y una oferta de contrato de mayor permanencia.')
    doc.add_heading('7. Actividad 6 — Evaluación de modelos', 1)
    doc.add_paragraph('Evalué ambos modelos sobre el mismo conjunto de prueba estratificado. Consideré Yes como la clase positiva, pues es la clase de clientes que se desea identificar.')
    mm,am=R['manual_metrics'],R['automatic_metrics']
    add_table(doc, ['Métrica','Red propuesta','Hill Climbing'], [['Accuracy',f"{mm['accuracy']:.4f}",f"{am['accuracy']:.4f}"],['Precision',f"{mm['precision']:.4f}",f"{am['precision']:.4f}"],['Recall',f"{mm['recall']:.4f}",f"{am['recall']:.4f}"],['F1-score',f"{mm['f1']:.4f}",f"{am['f1']:.4f}"]], [2.0,2.2,2.3])
    doc.add_paragraph(f'La red con Hill Climbing obtuvo mayor Accuracy ({am["accuracy"]:.2%}), Precision ({am["precision"]:.2%}) y F1-score ({am["f1"]:.2%}) que la estructura propuesta. Sin embargo, ambos modelos alcanzaron Recall de {am["recall"]:.2%}, por lo que detectan pocos abandonos reales. El desbalance de clases ayuda a explicar por qué una Accuracy relativamente alta no equivale a una detección adecuada de churn.')
    doc.add_heading('8. Actividad 7 — Comparación de enfoques', 1)
    add_table(doc, ['Aspecto','Red propuesta','Red Hill Climbing'], [['Origen de estructura','Conocimiento del problema','Aprendizaje desde los datos'],['Ventaja principal','Explicable y fácil de defender','Puede descubrir asociaciones no previstas'],['Limitación principal','Puede omitir relaciones reales','Depende de muestra, criterio y restricciones'],['Uso recomendado','Comunicación y validación de hipótesis','Exploración y punto de partida analítico']], [1.55,2.45,2.5])
    doc.add_paragraph('En mi análisis, la red automática produjo mejores métricas globales, pero la red propuesta es más directa para comunicar las hipótesis de negocio. No elegiría una sola sin más validación: usaría la propuesta como estructura interpretable y contrastaría sus relaciones con los hallazgos automáticos, pruebas temporales y conocimiento de expertos.')
    doc.add_heading('9. Conclusiones y recomendaciones', 1)
    add_bullet(doc, f'Identifiqué una tasa base de abandono de {R["churn_rate"]:.1%}; el escenario consultado eleva la estimación a {R["inference_probabilities"]["Yes"]:.2%}.')
    add_bullet(doc, 'La satisfacción baja, la falta de soporte técnico y la poca antigüedad son variables útiles para priorizar acciones de retención.')
    add_bullet(doc, 'La red Hill Climbing mejoró Accuracy, Precision y F1-score frente a la red propuesta, pero el Recall bajo indica que se requieren mejoras antes de desplegar el modelo.')
    add_bullet(doc, 'Como siguiente paso, probaría validación cruzada, ajuste del umbral de decisión, ponderación de la clase positiva y variables adicionales como interrupciones del servicio o historial de contacto.')
    doc.add_heading('Anexo — Reproducibilidad', 1)
    doc.add_paragraph('El archivo .ipynb adjunto contiene el flujo completo: carga de CSV, exploración, discretización, construcción de ambas redes, CPD completas, inferencia y evaluación. Para ejecutarlo en Google Colab, se debe subir el CSV al mismo entorno y ejecutar las celdas en orden. Se fija una semilla aleatoria (42) para mantener constante la partición entrenamiento/prueba.')
    doc.save(ROOT / 'Informe_Practica_1_Redes_Bayesianas_Adrian_Torrico.docx')

if __name__ == '__main__':
    create_notebook()
    create_docx()
    print('Entregables creados')
